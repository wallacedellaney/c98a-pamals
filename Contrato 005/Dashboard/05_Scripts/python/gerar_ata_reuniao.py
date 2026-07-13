"""
Gera a Ata de Reunião (RMA) mensal em .docx, a partir de:
  1. Áudio da gravação da reunião (Drive, pasta do Fechamento Mensal do mês) —
     transcrito localmente com Whisper (ver 00_Instrucoes/ata_reuniao.md).
  2. Planilha "RMA em andamento {MÊS}.xlsx" (mesma pasta) — abas 1.1 (horas
     voadas/faturadas), 1.2 (MMAM/Pontuação/IFD) e 4.1 (valores a faturar já
     com IFD aplicado). Fonte oficial, não recalculada por nós — evita a
     pequena divergência já conhecida entre o Cômputo Mensal automático do
     site e o número oficial da empresa (ver computo_mensal.md).
  3. Dados já tratados da nossa própria plataforma: Reparáveis, Empréstimos,
     Pagamentos, e o histórico de emergências (pra achar as aeronaves que
     negativaram no mês, a partir do Cômputo Mensal já calculado).
  4. O texto da Ata assinada de maio/2026 como modelo estrutural (seções,
     ordem, "Anexo A/B" no final em vez de tabela inteira no corpo).

O que é automático: indicadores (MMAM/P/PMAX/IFD), horas de voo e valores
(Módulo 1), notas fiscais, estatísticas de Reparáveis/Empréstimos, apuração
de entregas (IMR), e a lista de aeronaves com desempenho negativo (com o
detalhe de cada emergência, a partir da nossa base tratada).

O que NÃO é automático (a transcrição do áudio é só inserida ao final, em
"Transcrição da Reunião", pra revisão manual): abertura/objetivo em prosa,
discussão dos Módulos 2/3, MAPEM, encerramento, e a tabela de Pendências e
Encaminhamentos — são pontos de julgamento humano, não dá pra gerar prosa
formal confiável a partir de uma transcrição automática de áudio. O Word
final é pra editar (pedido do Wallace, 2026-07-13): o botão do site baixa
o .docx já com as seções automáticas prontas + a transcrição ao final.

Uso (teste manual, fora do site):
    python3 gerar_ata_reuniao.py 2026 6
"""

import io
import os
import sys
from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from PIL import Image, ImageDraw, ImageFont

from common import DADOS_TRATADOS, LOGS, registrar_log
from shared import drive_sync

MESES_PT = [
    "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
    "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro",
]
ABREV_MES = ["jan", "fev", "mar", "abr", "mai", "jun", "jul", "ago", "set", "out", "nov", "dez"]

# "Fechamentos mensais" (Drive) — pasta raiz, ano -> mês (ex.: "06 JUNHO").
PASTA_FECHAMENTOS_MENSAIS = "1PT5b2iqt2KVNmBjf1HDHPwcTLYHUNNjI"

PASTA_ATAS = DADOS_TRATADOS / "atas"

_MODELO_WHISPER = None  # cacheado em memória do processo — evita recarregar a cada chamada


def _fmt_moeda(valor):
    if valor is None or pd.isna(valor):
        return "—"
    texto = f"{valor:,.2f}"
    return "R$ " + texto.replace(",", "X").replace(".", ",").replace("X", ".")


def _fmt_num(valor, casas=2):
    """Formata número no padrão BR (vírgula decimal) — a Ata oficial usa
    vírgula em tudo (ex.: "90,46%", "0,95"), não só em moeda."""
    if valor is None or pd.isna(valor):
        return "—"
    return f"{valor:.{casas}f}".replace(".", ",")


def _fmt_data(valor):
    if valor is None or pd.isna(valor):
        return "—"
    return pd.Timestamp(valor).strftime("%d/%m/%Y")


def _fmt_horas_timedelta(td):
    if td is None or pd.isna(td):
        return "—"
    total_min = round(td.total_seconds() / 60)
    return f"{total_min // 60} horas e {total_min % 60} minutos"


# ---------------------------------------------------------------------------
# Drive: localizar a pasta do mês e os arquivos (áudio, RMA em andamento)
# ---------------------------------------------------------------------------

def _localizar_pasta_mes(ano, mes):
    anos = drive_sync.listar_pasta(PASTA_FECHAMENTOS_MENSAIS)
    pasta_ano = next((f for f in anos if f["name"].strip() == str(ano)), None)
    if pasta_ano is None:
        raise FileNotFoundError(f"Pasta do ano {ano} não encontrada em 'Fechamentos mensais' no Drive.")
    meses = drive_sync.listar_pasta(pasta_ano["id"])
    nome_mes = MESES_PT[mes - 1].upper()
    pasta_mes = next((f for f in meses if nome_mes in f["name"].upper()), None)
    if pasta_mes is None:
        raise FileNotFoundError(f"Pasta do mês {MESES_PT[mes - 1]}/{ano} não encontrada dentro de '{ano}' no Drive.")
    return drive_sync.listar_pasta(pasta_mes["id"])


def _baixar_audio_reuniao(arquivos_pasta):
    audio = next((f for f in arquivos_pasta if f["mimeType"].startswith("audio/")), None)
    if audio is None:
        return None, None
    conteudo = drive_sync.baixar_arquivo(audio["id"])
    return conteudo, audio["name"]


def _baixar_transcricao_manual(arquivos_pasta):
    """Transcrição escrita à mão pelo Wallace, salva no Drive (pasta do
    mês) com "transcri" no nome — preferida à transcrição automática do
    áudio (Whisper), pedido do Wallace em 2026-07-13 ("agora a
    transcrição eu vou colocar por escrito mesmo, no drive"). Suporta
    Google Doc nativo (exporta como texto) ou um .docx/.txt já enviado
    (baixa e extrai o texto)."""
    candidato = next((f for f in arquivos_pasta if "transcri" in f["name"].lower()), None)
    if candidato is None:
        return None

    if candidato["mimeType"] == "application/vnd.google-apps.document":
        conteudo = drive_sync.baixar_arquivo(candidato["id"], exportar_como=drive_sync.TEXTO_MIME)
        return conteudo.decode("utf-8") if isinstance(conteudo, bytes) else str(conteudo)

    conteudo = drive_sync.baixar_arquivo(candidato["id"])
    if candidato["name"].lower().endswith(".docx"):
        import io as _io
        from docx import Document as _Document
        doc = _Document(_io.BytesIO(conteudo))
        return "\n".join(p.text for p in doc.paragraphs)
    return conteudo.decode("utf-8") if isinstance(conteudo, bytes) else str(conteudo)


def _baixar_rma_em_andamento(arquivos_pasta):
    candidato = next(
        (f for f in arquivos_pasta if "rma em andamento" in f["name"].lower() and f["name"].lower().endswith(".xlsx")),
        None,
    )
    if candidato is None:
        raise FileNotFoundError("Arquivo 'RMA em andamento {MÊS}.xlsx' não encontrado na pasta do mês no Drive.")
    return drive_sync.baixar_arquivo(candidato["id"])


# ---------------------------------------------------------------------------
# Transcrição do áudio (Whisper local, cacheada em disco por mês)
# ---------------------------------------------------------------------------

def _preparar_ambiente_transcricao():
    """ffmpeg estático (via imageio-ffmpeg, sem depender de Homebrew/sudo) +
    certificados (via certifi, o Python.org deste Mac não vem com bundle
    próprio) — ver 00_Instrucoes/ata_reuniao.md."""
    import certifi
    import imageio_ffmpeg

    ffmpeg_dir = os.path.dirname(imageio_ffmpeg.get_ffmpeg_exe())
    if ffmpeg_dir not in os.environ.get("PATH", ""):
        os.environ["PATH"] = ffmpeg_dir + os.pathsep + os.environ.get("PATH", "")
    os.environ.setdefault("SSL_CERT_FILE", certifi.where())


def transcrever_audio(conteudo_bytes, forcar=False, caminho_cache=None):
    """Transcreve o áudio (bytes) em português com Whisper (modelo
    'medium' — testado em 2026-07-13, qualidade boa o bastante pra virar
    rascunho editável; o modelo 'base' errava números importantes). Cacheia
    em `caminho_cache` se informado, pra não reprocessar (~10 min) a cada
    clique no botão do site."""
    if caminho_cache and caminho_cache.exists() and not forcar:
        return caminho_cache.read_text(encoding="utf-8")

    _preparar_ambiente_transcricao()
    import tempfile

    import whisper

    global _MODELO_WHISPER
    if _MODELO_WHISPER is None:
        _MODELO_WHISPER = whisper.load_model("medium")

    with tempfile.NamedTemporaryFile(suffix=".m4a", delete=False) as tmp:
        tmp.write(conteudo_bytes)
        caminho_tmp = tmp.name
    try:
        resultado = _MODELO_WHISPER.transcribe(caminho_tmp, language="pt")
    finally:
        os.remove(caminho_tmp)

    texto = resultado["text"].strip()
    if caminho_cache:
        caminho_cache.parent.mkdir(parents=True, exist_ok=True)
        caminho_cache.write_text(texto, encoding="utf-8")
    return texto


# ---------------------------------------------------------------------------
# Parsers da planilha oficial "RMA em andamento {MÊS}.xlsx"
# ---------------------------------------------------------------------------

def _achar_linha(df, rotulo_contem):
    """Acha a linha que contém `rotulo_contem` (case-insensitive) em
    QUALQUER coluna — os rótulos da planilha não ficam sempre na coluna 0
    (ex.: "TOTAL SERV MOD 1" está na coluna 3). Mais resiliente do que
    fixar número de linha, já que a Tabela 1.1 cresce uma linha por mês.
    Usa a ÚLTIMA ocorrência: o título da aba (linha 0, tipo "Tabela 1.2 -
    ... (MMAM) e ... (IFD)") também contém o rótulo procurado, mas a linha
    de resumo com o valor de verdade sempre vem depois."""
    mascara = df.apply(lambda col: col.astype(str).str.contains(rotulo_contem, case=False, na=False)).any(axis=1)
    linhas = df[mascara]
    if linhas.empty:
        return None
    return linhas.iloc[-1]


def extrair_indicadores_rma(conteudo_bytes, ano, mes):
    """Lê as abas 1.1, 1.2 e 4.1 da planilha oficial e devolve um dict com
    os números que vão na Ata — nunca recalculado por nós (ver docstring do
    módulo)."""
    excel = pd.ExcelFile(io.BytesIO(conteudo_bytes))

    # --- 1.1: horas voadas/faturadas do mês de referência ---
    df_11 = pd.read_excel(excel, sheet_name="1.1", header=None)
    chave_mes = f"{ABREV_MES[mes - 1]}.-{str(ano)[2:]}"
    linha_11 = df_11[df_11[0].astype(str).str.strip() == chave_mes]
    if linha_11.empty:
        raise ValueError(f"Linha '{chave_mes}' não encontrada na aba 1.1 da planilha RMA em andamento.")
    linha_11 = linha_11.iloc[0]
    horas_voadas = linha_11[2]
    horas_faturadas = linha_11[4]

    # --- 1.2: MMAM / Pontuação / IFD ---
    df_12 = pd.read_excel(excel, sheet_name="1.2", header=None)
    mmam = _achar_linha(df_12, "MMAM")[1]
    pontuacao_obtida = _achar_linha(df_12, "PONTUAÇÃO OBTIDA")[1]
    pont_max = _achar_linha(df_12, "PONT. MAX")[1]
    ifd = _achar_linha(df_12, "Índice Final de Desempenho")[1]

    # --- 4.1: valores a faturar (Módulo 1), já com IFD aplicado ---
    df_41 = pd.read_excel(excel, sheet_name="4.1", header=None)
    linha_mod1 = df_41.iloc[6]  # primeira linha de dado do Módulo 1 (Serviço de Gerenciamento Logístico)
    valor_unitario_hora = linha_mod1[3]
    valor_total_antes_ifd = _achar_linha(df_41, "SUBTOTAL SERV MOD 1")[7]
    valor_total_modulo1 = _achar_linha(df_41, "TOTAL SERV MOD 1")[9]
    linha_multa = _achar_linha(df_41, "Multa sobre itens entreg")
    multa = linha_multa[9] if linha_multa is not None else 0
    valor_total_mes = _achar_linha(df_41, "VALOR TOTAL A SER FATURADO")[9]

    return {
        "horas_voadas": horas_voadas,
        "horas_faturadas": horas_faturadas,
        "mmam": float(mmam),
        "pontuacao_obtida": float(pontuacao_obtida),
        "pont_max": float(pont_max),
        "ifd": float(ifd),
        "valor_unitario_hora": float(valor_unitario_hora),
        "valor_total_antes_ifd": float(valor_total_antes_ifd),
        "valor_total_modulo1": float(valor_total_modulo1),
        "multa": float(multa) if pd.notna(multa) else 0.0,
        "valor_total_mes": float(valor_total_mes),
    }


# ---------------------------------------------------------------------------
# Dados da nossa própria plataforma (Reparáveis, Empréstimos, Pagamentos,
# aeronaves com desempenho negativo, apuração de entregas)
# ---------------------------------------------------------------------------

def carregar_reparaveis_abertas():
    caminho = DADOS_TRATADOS / "base_reparaveis_tratada.xlsx"
    df = pd.read_excel(caminho)
    return df[df["em_aberto"]].copy()


def carregar_emprestimos_mes(ano, mes):
    caminho = DADOS_TRATADOS / "base_devolucoes_tratada.xlsx"
    df = pd.read_excel(caminho)
    df["pedido_envio"] = pd.to_datetime(df["pedido_envio"], errors="coerce")
    df["quantidade_efetiva"] = df["quantidade"].fillna(1)
    return df[
        (df["pedido_envio"].dt.year == ano) & (df["pedido_envio"].dt.month == mes)
    ].copy()


def carregar_pagamentos():
    caminho = DADOS_TRATADOS / "base_pagamentos_tratada.xlsx"
    df_pagamentos = pd.read_excel(caminho, sheet_name="Pagamentos")
    df_contrato = pd.read_excel(caminho, sheet_name="Contrato")
    return df_pagamentos, df_contrato.iloc[0].to_dict()


def carregar_aeronaves_negativadas(ano, mes):
    """Junta o Cômputo Mensal (motivos de negativação, já calculado pelo
    site) com o histórico completo de emergências, pra ter o detalhe rico
    (PN, nomenclatura, prazo, DPE, atraso, AWB, observação) de cada
    emergência que derrubou alguma aeronave no mês."""
    pasta = DADOS_TRATADOS / "computo_mensal"
    caminho_motivos = pasta / f"{ano}-{mes:02d}_motivos.csv"
    if not caminho_motivos.exists():
        return {}
    df_motivos = pd.read_csv(caminho_motivos, dtype={"matricula": str})
    if df_motivos.empty:
        return {}

    caminho_hist = DADOS_TRATADOS / "historico_completo_emergencias.xlsx"
    df_hist = pd.read_excel(caminho_hist, dtype={"matricula_aeronave": str})
    df_hist["numero_emergencia"] = df_hist["numero_emergencia"].astype(str)
    df_motivos["numero_emergencia"] = df_motivos["numero_emergencia"].astype(str)

    detalhes = df_motivos.merge(
        df_hist, on="numero_emergencia", how="left", suffixes=("", "_hist")
    )

    aeronaves = {}
    for matricula, grupo in detalhes.groupby("matricula"):
        aeronaves[matricula] = grupo.to_dict("records")
    return aeronaves


def carregar_apuracao_entregas(ano, mes):
    """Mesma regra da aba 'Atrasos' do site (fechamento_mensal.py::_atrasos,
    ponto 2): tudo que foi concluído/cancelado dentro do mês de referência,
    não importa quando abriu."""
    caminho_hist = DADOS_TRATADOS / "historico_completo_emergencias.xlsx"
    df = pd.read_excel(caminho_hist)
    df["atendido_cancelado_dt"] = pd.to_datetime(df["atendido_cancelado"], errors="coerce")
    inicio_mes = pd.Timestamp(ano, mes, 1)
    fim_mes = inicio_mes + pd.offsets.MonthEnd(0)
    concluidas = df[
        (~df["em_aberto"])
        & (df["atendido_cancelado_dt"] >= inicio_mes)
        & (df["atendido_cancelado_dt"] <= fim_mes)
    ]
    total = len(concluidas)
    no_prazo = int((concluidas["dias_atraso"] <= 0).sum())
    pct = (100 * no_prazo / total) if total else 0.0
    return {"total_previstas": total, "no_prazo": no_prazo, "atrasadas": total - no_prazo, "pct": pct}


# ---------------------------------------------------------------------------
# Montagem do .docx
# ---------------------------------------------------------------------------

AZUL_TITULO = RGBColor(0x1F, 0x3A, 0x5F)


def _titulo(doc, texto, nivel=1):
    p = doc.add_heading(texto, level=nivel)
    for run in p.runs:
        run.font.color.rgb = AZUL_TITULO
    return p


def _paragrafo(doc, texto, negrito=False, centralizado=False):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = negrito
    if centralizado:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _tabela(doc, cabecalho, linhas):
    tabela = doc.add_table(rows=1, cols=len(cabecalho))
    tabela.style = "Table Grid"
    for i, titulo_col in enumerate(cabecalho):
        celula = tabela.rows[0].cells[i]
        celula.text = str(titulo_col)
        for p in celula.paragraphs:
            for run in p.runs:
                run.bold = True
    for linha in linhas:
        celulas = tabela.add_row().cells
        for i, valor in enumerate(linha):
            celulas[i].text = str(valor)
    return tabela


FONTE_TABELA_IMAGEM = "/System/Library/Fonts/Supplemental/Arial.ttf"
COR_CABECALHO_IMAGEM = (232, 232, 232)
COR_BORDA_IMAGEM = (170, 170, 170)


def _renderizar_tabela_imagem(df, colunas, caminho_png, largura_px=2000, altura_linha=32, tamanho_fonte=16):
    """Desenha `df[colunas]` como uma imagem de tabela (não uma tabela
    nativa do Word) — pedido do Wallace em 2026-07-13 ("queria tipo
    imagem para caber melhor igual ta no exemplo"): a Ata assinada de
    maio também traz o Anexo A como imagem (captura de tela), não como
    tabela editável. Larguras de coluna proporcionais ao maior texto de
    cada coluna (cabeçalho + dados), sem depender de nenhuma lib pesada
    (só Pillow, que já vem instalado)."""
    try:
        fonte = ImageFont.truetype(FONTE_TABELA_IMAGEM, tamanho_fonte)
        fonte_negrito = ImageFont.truetype(FONTE_TABELA_IMAGEM.replace("Arial.ttf", "Arial Bold.ttf"), tamanho_fonte)
    except OSError:
        fonte = fonte_negrito = ImageFont.load_default()

    cabecalhos = colunas
    dados = [[_formatar_valor_generico(registro.get(c)) for c in colunas] for _, registro in df.iterrows()]

    img_medida = Image.new("RGB", (10, 10))
    desenho_medida = ImageDraw.Draw(img_medida)

    def _largura_texto(texto, fnt):
        return desenho_medida.textbbox((0, 0), texto, font=fnt)[2]

    larguras = []
    for i, col in enumerate(cabecalhos):
        maior = _largura_texto(str(col), fonte_negrito)
        for linha in dados:
            maior = max(maior, _largura_texto(linha[i], fonte))
        larguras.append(maior + 24)

    largura_total_conteudo = sum(larguras)
    escala = largura_px / largura_total_conteudo if largura_total_conteudo else 1
    larguras = [int(w * escala) for w in larguras]
    largura_px = sum(larguras)
    altura_px = altura_linha * (len(dados) + 1)

    img = Image.new("RGB", (largura_px, altura_px), "white")
    desenho = ImageDraw.Draw(img)

    def _linha(y, valores, fnt, cor_fundo):
        if cor_fundo:
            desenho.rectangle([0, y, largura_px, y + altura_linha], fill=cor_fundo)
        x = 0
        for i, valor in enumerate(valores):
            desenho.text((x + 6, y + (altura_linha - tamanho_fonte) // 2 - 2), valor, font=fnt, fill=(20, 20, 20))
            x += larguras[i]
        desenho.line([(0, y), (largura_px, y)], fill=COR_BORDA_IMAGEM, width=1)

    _linha(0, [str(c) for c in cabecalhos], fonte_negrito, COR_CABECALHO_IMAGEM)
    for i, linha in enumerate(dados, start=1):
        _linha(i * altura_linha, linha, fonte, None)
    desenho.line([(0, altura_px - 1), (largura_px, altura_px - 1)], fill=COR_BORDA_IMAGEM, width=1)
    x = 0
    for w in larguras:
        desenho.line([(x, 0), (x, altura_px)], fill=COR_BORDA_IMAGEM, width=1)
        x += w
    desenho.line([(largura_px - 1, 0), (largura_px - 1, altura_px)], fill=COR_BORDA_IMAGEM, width=1)

    caminho_png.parent.mkdir(parents=True, exist_ok=True)
    img.save(caminho_png)
    return caminho_png, largura_px, altura_px


def _formatar_valor_generico(valor):
    if valor is None or (isinstance(valor, (float, pd.Timestamp, type(pd.NaT))) and pd.isna(valor)):
        return ""
    if isinstance(valor, (pd.Timestamp, datetime)):
        return valor.strftime("%d/%m/%Y")
    if isinstance(valor, float) and valor.is_integer():
        return str(int(valor))
    return str(valor)


def _texto_revisar(doc, rotulo):
    p = doc.add_paragraph()
    run = p.add_run(f"[{rotulo} — revisar a partir da transcrição da reunião, ao final deste documento]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)


def _bullet_emergencia(doc, registro):
    numero = registro.get("numero_emergencia", "—")
    pn = registro.get("pn", "—")
    nomenclatura = registro.get("nomenclatura", "—")
    categoria = registro.get("categoria", "—")
    abertura = _fmt_data(registro.get("data_abertura"))
    info = _fmt_data(registro.get("data_info"))
    prazo = _fmt_data(registro.get("prazo_entrega"))
    dpe = registro.get("dpe")
    dpe_txt = _fmt_data(dpe) if dpe and not pd.isna(dpe) else "—"
    dias_atraso = registro.get("dias_atraso")
    if pd.notna(dias_atraso):
        if dias_atraso > 0:
            situacao_txt = f"Atraso: {int(dias_atraso)} dia(s)."
        elif dias_atraso < 0:
            situacao_txt = f"Antecipação: {int(-dias_atraso)} dia(s)."
        else:
            situacao_txt = "Sem atraso."
    else:
        situacao_txt = "Situação de prazo não apurada."

    partes = [
        f"Emg nº {numero} – PN {pn} ({nomenclatura}, Cat. {categoria}): "
        f"Emg iniciada em {abertura}, informada em {info}, prazo de entrega {prazo}, "
        f"DPE VEE ONE {dpe_txt}. {situacao_txt}"
    ]
    awb = registro.get("awb")
    if awb and pd.notna(awb):
        partes.append(f"AWB {awb}.")
    obs = registro.get("obs_coordenadoria_fiscal") or registro.get("obs_vee_one")
    if obs and pd.notna(obs):
        partes.append(str(obs))

    p = doc.add_paragraph(style="List Bullet")
    p.add_run(" ".join(partes))


def gerar_ata(ano, mes, caminho_saida, forcar_transcricao=False):
    arquivos_lidos = []
    arquivos_pasta = _localizar_pasta_mes(ano, mes)

    conteudo_rma = _baixar_rma_em_andamento(arquivos_pasta)
    arquivos_lidos.append(f"Drive: RMA em andamento {MESES_PT[mes-1]}.xlsx")
    indicadores = extrair_indicadores_rma(conteudo_rma, ano, mes)

    # Transcrição manual (escrita pelo Wallace, salva no Drive) tem
    # prioridade sobre a automática por áudio — pedido do Wallace em
    # 2026-07-13. Só cai pro áudio+Whisper se não achar nenhum arquivo de
    # transcrição na pasta do mês.
    transcricao = _baixar_transcricao_manual(arquivos_pasta) or ""
    if transcricao:
        arquivos_lidos.append("Drive: transcrição (escrita manualmente)")
    else:
        conteudo_audio, nome_audio = _baixar_audio_reuniao(arquivos_pasta)
        if conteudo_audio:
            arquivos_lidos.append(f"Drive: {nome_audio}")
            caminho_cache = PASTA_ATAS / f"transcricao_{ano}-{mes:02d}.txt"
            transcricao = transcrever_audio(conteudo_audio, forcar=forcar_transcricao, caminho_cache=caminho_cache)

    reparaveis = carregar_reparaveis_abertas()
    arquivos_lidos.append("base_reparaveis_tratada.xlsx")
    emprestimos_mes = carregar_emprestimos_mes(ano, mes)
    arquivos_lidos.append("base_devolucoes_tratada.xlsx")
    df_pagamentos, resumo_contrato = carregar_pagamentos()
    arquivos_lidos.append("base_pagamentos_tratada.xlsx")
    aeronaves_negativadas = carregar_aeronaves_negativadas(ano, mes)
    arquivos_lidos.append(f"computo_mensal/{ano}-{mes:02d}_motivos.csv + historico_completo_emergencias.xlsx")
    apuracao = carregar_apuracao_entregas(ano, mes)

    mes_nome = MESES_PT[mes - 1]
    doc = Document()

    _paragrafo(doc, "MINISTÉRIO DA DEFESA", centralizado=True)
    _paragrafo(doc, "COMANDO DA AERONÁUTICA", centralizado=True)
    _paragrafo(doc, "PARQUE DE MATERIAL AERONÁUTICO DE LAGOA SANTA", centralizado=True)
    _paragrafo(doc, "DIVISÃO TÉCNICA", centralizado=True)
    doc.add_paragraph()
    _titulo(doc, "Reunião Mensal de Acompanhamento", nivel=0)
    _titulo(doc, "Ata de Reunião", nivel=1)
    _paragrafo(doc, f"{mes_nome} de {ano}", centralizado=True)

    _titulo(doc, "Participantes")
    _paragrafo(
        doc,
        f"Reuniram-se, em data referente ao mês de {mes_nome.lower()} de {ano}, representantes do Parque de "
        "Material Aeronáutico de Lagoa Santa (PAMALS) e da empresa Vee One Manutenção e Serviços Técnicos LTDA "
        "(contratada), para tratarem de assuntos relacionados às atividades de Suporte Logístico da frota de "
        "aeronaves C-98 da Força Aérea Brasileira, no âmbito do Contrato nº 005/CELOG-PAMALS/2025, tendo por "
        f"referência os resultados apurados no mês de {mes_nome.lower()} de {ano}."
    )
    _texto_revisar(doc, "Lista de participantes (PAMALS e Contratada)")

    _titulo(doc, "Prestação de Contas no Mês de Referência")
    _paragrafo(
        doc,
        "Os valores apresentados a seguir já contemplam a aplicação do Índice Final de Desempenho (IFD), "
        f"conforme memória de cálculo constante na planilha da RMA {mes:02d}/{ano} – Módulo 1."
    )
    _tabela(doc, ["Descrição", "Valor"], [
        ["Módulo 1 Faturamento – Remunerado por Hora de Voo", _fmt_moeda(indicadores["valor_total_modulo1"])],
        ["Módulo 2 Faturamento – Orçamentos de Material (conforme orçamentos em aberto)", "—"],
        ["Módulo 3 Faturamento – GPS / Manuais (conforme orçamentos em aberto)", "—"],
    ])
    _paragrafo(
        doc,
        f"No que se refere ao Módulo 1, o valor discriminado da hora de voo corresponde a "
        f"{_fmt_moeda(indicadores['valor_unitario_hora'])}, calculado com base nas horas efetivamente realizadas "
        f"no período, totalizando {_fmt_horas_timedelta(indicadores['horas_voadas'])}. O valor bruto apurado foi de "
        f"{_fmt_moeda(indicadores['valor_total_antes_ifd'])}. Aplicado o Índice Final de Desempenho (IFD) de "
        f"{_fmt_num(indicadores['ifd'], 3)}, o valor final apurado para o Módulo 1 foi de "
        f"{_fmt_moeda(indicadores['valor_total_modulo1'])}, com multa sobre itens entregues em atraso de "
        f"{_fmt_moeda(indicadores['multa'])}."
    )

    _titulo(doc, "Relatos da Reunião")

    _titulo(doc, "1. Abertura e Objetivo", nivel=2)
    _paragrafo(
        doc,
        f"O Fiscal do Contrato deu início à reunião com a apresentação dos objetivos da RMA de {mes_nome.lower()}, "
        f"esclarecendo que a pauta contemplaria a média mensal de aeronaves montadas do Contrato "
        f"nº 005/CELOG-PAMALS/2025, referente ao mês de {mes_nome.lower()} de {ano}, bem como demais informações "
        "pertinentes à execução contratual."
    )

    _titulo(doc, "2. Visão Macro do Contrato e Indicadores de Desempenho", nivel=2)
    _paragrafo(
        doc,
        f"Foi apresentada a visão macro do contrato, com ênfase nos indicadores de desempenho do período. A "
        f"Tabela 1.2 (Média Mensal de Aeronaves Montadas – MMAM e Índice Final de Desempenho – IFD) foi exibida, "
        f"contendo os registros diários das aeronaves da frota ao longo do mês de {mes_nome.lower()} de {ano}. Os "
        "resultados consolidados foram os seguintes:"
    )
    _tabela(doc, ["Indicador", "Valor"], [
        ["Média Mensal de Aeronaves Montadas (MMAM)", f"{_fmt_num(indicadores['mmam'] * 100)}%"],
        ["Pontuação Obtida (P)", f"{indicadores['pontuacao_obtida']:.0f}"],
        ["Pontuação Máxima no Mês (PMAX)", f"{indicadores['pont_max']:.0f}"],
        ["Índice Final de Desempenho (IFD)", _fmt_num(indicadores['ifd'])],
    ])

    _titulo(doc, "3. Aeronaves com Desempenho Negativo – Emergências", nivel=2)
    if aeronaves_negativadas:
        _paragrafo(
            doc,
            "Na sequência, foram apresentadas as aeronaves que apresentaram desempenho negativo no período, com "
            "detalhamento das emergências registradas por matrícula, conforme abaixo:"
        )
        for matricula in sorted(aeronaves_negativadas):
            _paragrafo(doc, f"Aeronave FAB {matricula}", negrito=True)
            for registro in aeronaves_negativadas[matricula]:
                _bullet_emergencia(doc, registro)
    else:
        _paragrafo(doc, "Nenhuma aeronave apresentou desempenho negativo (desmontada) no período.")

    _titulo(doc, "4. Horas de Voo e Valor Contratual", nivel=2)
    _paragrafo(
        doc,
        f"Foi apresentado o extrato de horas de voo do período, totalizando "
        f"{_fmt_horas_timedelta(indicadores['horas_voadas'])} voadas e "
        f"{_fmt_horas_timedelta(indicadores['horas_faturadas'])} faturadas. O valor unitário da hora de voo "
        f"utilizado foi de {_fmt_moeda(indicadores['valor_unitario_hora'])}, resultando em um valor bruto de "
        f"{_fmt_moeda(indicadores['valor_total_antes_ifd'])}. Aplicado o IFD de {_fmt_num(indicadores['ifd'], 3)}, o valor "
        f"final do Módulo 1 apurado foi de {_fmt_moeda(indicadores['valor_total_modulo1'])}, com multa por itens "
        f"entregues em atraso de {_fmt_moeda(indicadores['multa'])}."
    )

    _titulo(doc, "5. Notas Fiscais e Status de Pagamento", nivel=2)
    pendentes = df_pagamentos[df_pagamentos["pendente"].notna() & (df_pagamentos["pendente"] > 0)]
    _paragrafo(
        doc,
        f"Situação geral do contrato: valor total pendente de faturamento de "
        f"{_fmt_moeda(resumo_contrato.get('saldo_a_faturar'))}. "
        + (f"{len(pendentes)} nota(s) fiscal(is) com valor pendente de pagamento:" if len(pendentes) else
           "Não há notas fiscais com valor pendente no momento.")
    )
    if len(pendentes):
        _tabela(doc, ["Referência", "NF", "Valor", "Pendente", "Situação"], [
            [r["referencia"], r["numero_nota_fiscal"], _fmt_moeda(r["valor_nfs"]), _fmt_moeda(r["pendente"]), r["situacao"]]
            for _, r in pendentes.iterrows()
        ])

    _titulo(doc, "6. Módulos 2 e 3 – Orçamentos", nivel=2)
    _texto_revisar(doc, "Discussão sobre orçamentos dos Módulos 2 e 3")

    _titulo(doc, "7. Controle de Empréstimos e Devoluções", nivel=2)
    total_qtd = emprestimos_mes["quantidade_efetiva"].sum()
    _paragrafo(
        doc,
        f"Foi apresentado o controle de empréstimos de materiais realizados no mês de {mes_nome.lower()} de {ano}: "
        f"{len(emprestimos_mes)} linha(s) de pedido, totalizando {total_qtd:,.0f} de quantidade (soma da coluna "
        "Quantidade de cada linha). O detalhamento de PN, nomenclatura, quantidade, pedido/emergência, motivo, "
        "aeronave atendida e destino consta no Anexo A desta ata."
    )

    _titulo(doc, "8. Controle de Reparáveis", nivel=2)
    _paragrafo(
        doc,
        f"Foi retomado o assunto do controle de reparáveis: {len(reparaveis)} Ordem(ns) de Serviço em aberto no "
        "momento da extração. O controle completo consta no Anexo B desta ata."
    )

    _titulo(doc, "9. Apuração de Entregas – IMR", nivel=2)
    _paragrafo(
        doc,
        f"Foram apresentados os dados da apuração de entregas do período: de {apuracao['total_previstas']} "
        f"entregas previstas, {apuracao['no_prazo']} foram realizadas no prazo, resultando em um índice de "
        f"{_fmt_num(apuracao['pct'])}%."
    )

    _titulo(doc, "10. Controle do MAPEM e devoluções de itens emprestados", nivel=2)
    _texto_revisar(doc, "Discussão sobre MAPEM e devoluções")

    _titulo(doc, "11. Encerramento", nivel=2)
    _texto_revisar(doc, "Observações de encerramento")

    _titulo(doc, "Pendências e Encaminhamentos")
    _tabela(doc, ["Nº", "Pendência / Encaminhamento", "Responsável", "Prazo"], [
        ["1", "[revisar a partir da transcrição da reunião]", "—", "—"],
    ])

    doc.add_page_break()
    _titulo(doc, "Assinaturas")
    _paragrafo(doc, "Representante do PAMALS:")
    _paragrafo(doc, "_" * 40)
    _paragrafo(doc, "Fiscal do Contrato")
    doc.add_paragraph()
    _paragrafo(doc, "Representante da Contratada:")
    _paragrafo(doc, "_" * 40)
    _paragrafo(doc, "Preposto – VEE-ONE Manutenção e Serviços Técnicos LTDA")

    doc.add_page_break()
    _titulo(doc, "ANEXO A — Controle de Devoluções (Empréstimos do mês)")
    _paragrafo(doc, f"Contrato nº 005/CELOG-PAMALS/2025 — {mes_nome} de {ano}", negrito=True)
    colunas_emp = [
        "numero_ordem", "part_number", "descricao", "quantidade_texto", "pedido_emg",
        "motivo", "anv", "destino", "status",
    ]
    caminho_imagem_emp = caminho_saida.parent / f"_anexo_a_emprestimos_{ano}-{mes:02d}.png"
    _renderizar_tabela_imagem(emprestimos_mes, colunas_emp, caminho_imagem_emp)
    doc.add_picture(str(caminho_imagem_emp), width=Inches(6.5))
    caminho_imagem_emp.unlink()  # já embutida no .docx, não precisa sobrar no disco

    doc.add_page_break()
    _titulo(doc, "ANEXO B — Controle de Reparáveis (C-98, em aberto)")
    _paragrafo(doc, f"Contrato nº 005/CELOG-PAMALS/2025 — extraído em {datetime.now().strftime('%d/%m/%Y')}", negrito=True)
    colunas_rep = ["os", "pn", "nomenclatura", "unidade_solicitante", "situacao", "onde_se_encontra", "tat_siloms"]
    _tabela(doc, colunas_rep, reparaveis[colunas_rep].fillna("—").values.tolist())

    if transcricao:
        doc.add_page_break()
        _titulo(doc, "Transcrição da Reunião (rascunho automático — revisar)")
        _paragrafo(
            doc,
            "Transcrição gerada automaticamente a partir do áudio da reunião (Whisper, local). Pode conter erros "
            "de nomes/termos técnicos — usar como referência pra completar as seções marcadas acima, não como "
            "texto final."
        )
        _paragrafo(doc, transcricao)

    caminho_saida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(caminho_saida))

    registrar_log(
        f"gerar_ata_reuniao_{ano}-{mes:02d}",
        arquivos_lidos,
        [str(caminho_saida)],
        [],
        proximas_acoes=["Revisar seções marcadas '[revisar a partir da transcrição]' e a tabela de Pendências antes de assinar."],
    )
    return caminho_saida


if __name__ == "__main__":
    ano_arg = int(sys.argv[1])
    mes_arg = int(sys.argv[2])
    saida = DADOS_TRATADOS / "atas" / f"Ata_RMA_{MESES_PT[mes_arg-1]}_{ano_arg}_rascunho.docx"
    resultado = gerar_ata(ano_arg, mes_arg, saida)
    print(f"Gerado: {resultado}")
